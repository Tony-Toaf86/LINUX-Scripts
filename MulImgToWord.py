from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# Nombre de la Carpeta donde están tus imágenes
folder = "img" 

# Crear documento
doc = Document()
doc.add_heading("Colección de imágenes", level=1)

# Listar y ordenar las imágenes
imagenes = sorted([f for f in os.listdir(folder) if f.lower().endswith((".jpg", ".jpeg", ".png"))])

# Contador
count = 0

for filename in imagenes:
    image_path = os.path.join(folder, filename)

    # Nombre de la imagen (centrado)
    p = doc.add_paragraph(filename)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(10)

    # Imagen (centrada y tamaño reducido para que quepan 3 en una página)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(2.4))  # <-- AJUSTE DE TAMAÑO

    count += 1

    # Cada 3 imágenes → salto de página
    if count % 3 == 0:
        doc.add_page_break()

# Guardar documento
doc.save("imagenes.docx")
print("Documento creado con éxito: imagenes.docx")
